from __future__ import annotations

import csv
import math
import os
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "report"
FIG_DIR = OUT_DIR / "figures"
CSV_PATH = ROOT / "backend" / "data" / "PulseBoard_APM_14400_Dataset_With_HTTP_Status.csv"
REFERENCE_PAGE = ROOT / "tmp" / "reference_pdf" / "page-01.png"
DOCX_PATH = OUT_DIR / "PulseBoard_APM_Design_Project_Report.docx"

GREEN = "008000"
BLACK = "000000"
LIGHT_GREY = "F2F2F2"
MID_GREY = "D9E1F2"
SOFT_GREEN = "E2F0D9"
SOFT_YELLOW = "FFF2CC"
SOFT_RED = "FCE4D6"
NAVY = "17365D"
MUTED = "666666"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def load_rows() -> list[dict[str, str]]:
    with CSV_PATH.open(encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def mean(rows: list[dict[str, str]], field: str) -> float:
    values = [float(row[field]) for row in rows]
    return sum(values) / len(values)


def percentile(values: list[float], fraction: float) -> float:
    ordered = sorted(values)
    index = min(len(ordered) - 1, math.floor(len(ordered) * fraction))
    return ordered[index]


def calculate_stats(rows: list[dict[str, str]]) -> dict:
    by_app: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        by_app[row["Application"]].append(row)

    applications = []
    for app, app_rows in by_app.items():
        applications.append(
            {
                "application": app,
                "response": mean(app_rows, "Response_Time_ms"),
                "cpu": mean(app_rows, "CPU_Usage_%"),
                "memory": mean(app_rows, "Memory_Usage_%"),
                "availability": mean(app_rows, "Availability_%"),
                "error": mean(app_rows, "Error_Rate_%"),
                "throughput": mean(app_rows, "Request_Throughput_RPM"),
                "p95": percentile([float(r["P95_Latency_ms"]) for r in app_rows], 0.95),
            }
        )
    applications.sort(key=lambda item: item["response"], reverse=True)

    return {
        "records": len(rows),
        "applications": len(by_app),
        "regions": Counter(row["Region"] for row in rows),
        "environments": Counter(row["Environment"] for row in rows),
        "severity": Counter(row["Severity"] for row in rows),
        "events": Counter(row["Event"] for row in rows),
        "start": min(row["Timestamp"] for row in rows),
        "end": max(row["Timestamp"] for row in rows),
        "response": mean(rows, "Response_Time_ms"),
        "cpu": mean(rows, "CPU_Usage_%"),
        "memory": mean(rows, "Memory_Usage_%"),
        "availability": mean(rows, "Availability_%"),
        "error": mean(rows, "Error_Rate_%"),
        "throughput": mean(rows, "Request_Throughput_RPM"),
        "http4xx": sum(int(row["HTTP_4XX_Count"]) for row in rows),
        "http5xx": sum(int(row["HTTP_5XX_Count"]) for row in rows),
        "p95": percentile([float(row["P95_Latency_ms"]) for row in rows], 0.95),
        "applications_data": applications,
    }


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def center_text(draw, xy, text: str, font_obj, fill="#1F1F1F") -> None:
    x, y = xy
    bounds = draw.multiline_textbbox((0, 0), text, font=font_obj, spacing=5, align="center")
    draw.multiline_text((x - (bounds[2] - bounds[0]) / 2, y - (bounds[3] - bounds[1]) / 2), text, font=font_obj, fill=fill, spacing=5, align="center")


def rounded_box(draw, coords, fill, outline="#4F81BD", width=3, radius=16) -> None:
    draw.rounded_rectangle(coords, radius=radius, fill=fill, outline=outline, width=width)


def save_figures(stats: dict) -> dict[str, Path]:
    """Create report figures using Pillow, which is part of the bundled workspace runtime."""
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    title_font = font(40, True)
    subtitle_font = font(26, True)
    text_font = font(22)
    small_font = font(19)

    # Figure 1: architecture
    image = Image.new("RGB", (2040, 960), "white")
    draw = ImageDraw.Draw(image)
    center_text(draw, (1020, 100), "PulseBoard APM - Logical Architecture", title_font, "#17365D")
    boxes = [
        (80, 390, 430, 620, "APM CSV Dataset\n14,400 records", "#DDEBF7"),
        (570, 390, 920, 620, "Node.js + Express\nREST API", "#E2F0D9"),
        (1060, 390, 1410, 620, "Metrics, Health\nand Alert Services", "#FFF2CC"),
        (1550, 390, 1900, 620, "React + Recharts\nDashboard", "#FCE4D6"),
    ]
    for x1, y1, x2, y2, label, fill in boxes:
        rounded_box(draw, (x1, y1, x2, y2), fill)
        center_text(draw, ((x1 + x2) / 2, (y1 + y2) / 2), label, subtitle_font)
    for x1, x2 in [(440, 560), (930, 1050), (1420, 1540)]:
        draw.line((x1, 505, x2 - 24, 505), fill="#4F81BD", width=6)
        draw.polygon([(x2 - 24, 490), (x2 - 24, 520), (x2, 505)], fill="#4F81BD")
    center_text(draw, (1020, 790), "Filters select application, region and time range; the API computes metrics and returns dashboard-ready JSON.", text_font, "#555555")
    architecture = FIG_DIR / "architecture.png"
    image.save(architecture)

    # Figure 2: composition
    image = Image.new("RGB", (2040, 900), "white")
    draw = ImageDraw.Draw(image)
    center_text(draw, (1020, 62), "Figure 2. Dataset Composition", title_font, "#17365D")
    center_text(draw, (510, 150), "Event Severity Composition", subtitle_font)
    centre_x = 510
    base_y, max_height = 760, 470
    sev_order = ["Info", "Warning", "Critical"]
    sev_values = [stats["severity"].get(label, 0) for label in sev_order]
    colours = ["#5B9BD5", "#FFC000", "#C00000"]
    for index, (label, value, colour) in enumerate(zip(sev_order, sev_values, colours)):
        x1 = 175 + index * 220
        height = int(max_height * value / max(sev_values))
        draw.rectangle((x1, base_y - height, x1 + 110, base_y), fill=colour)
        center_text(draw, (x1 + 55, base_y - height - 35), f"{value:,}", small_font)
        center_text(draw, (x1 + 55, base_y + 42), label, small_font)
    draw.line((125, base_y, 850, base_y), fill="#555555", width=2)
    draw.line((125, 230, 125, base_y), fill="#555555", width=2)
    center_text(draw, (1530, 150), "Regional Data Coverage", subtitle_font)
    regions = list(stats["regions"].keys())
    max_region = max(stats["regions"].values())
    for index, region in enumerate(regions):
        y = 310 + index * 135
        value = stats["regions"][region]
        width = int(550 * value / max_region)
        draw.text((1120, y + 15), region, font=text_font, fill="#1F1F1F")
        draw.rounded_rectangle((1350, y, 1350 + width, y + 55), radius=8, fill="#70AD47")
        draw.text((1360 + width, y + 14), f"{value:,}", font=small_font, fill="#1F1F1F")
    draw.line((1350, 730, 1900, 730), fill="#555555", width=2)
    composition = FIG_DIR / "dataset_composition.png"
    image.save(composition)

    # Figure 3: application performance
    image = Image.new("RGB", (2040, 1420), "white")
    draw = ImageDraw.Draw(image)
    center_text(draw, (1020, 55), "Figure 3. Application-Level Performance", title_font, "#17365D")
    center_text(draw, (1020, 125), "Average Response Time by Application", subtitle_font)
    apps = stats["applications_data"]
    response_left, response_right = 560, 1740
    chart_top, row_height = 185, 65
    max_response = 450
    threshold_x = response_left + (300 / max_response) * (response_right - response_left)
    draw.line((threshold_x, chart_top - 15, threshold_x, chart_top + len(apps) * row_height), fill="#C00000", width=4)
    draw.text((threshold_x + 8, chart_top - 46), "Warning threshold (300 ms)", font=small_font, fill="#C00000")
    for index, app in enumerate(apps):
        y = chart_top + index * row_height
        draw.text((80, y + 13), app["application"], font=small_font, fill="#1F1F1F")
        width = int((app["response"] / max_response) * (response_right - response_left))
        draw.rounded_rectangle((response_left, y, response_left + width, y + 35), radius=6, fill="#4472C4")
        draw.text((response_left + width + 12, y + 6), f"{app['response']:.1f} ms", font=small_font, fill="#1F1F1F")
    draw.line((response_left, chart_top + len(apps) * row_height + 5, response_right, chart_top + len(apps) * row_height + 5), fill="#555555", width=2)
    center_text(draw, (1020, 920), "Average Availability by Application", subtitle_font)
    availability_left, availability_right = 560, 1740
    avail_top, avail_low, avail_high = 980, 99.55, 99.85
    for index, app in enumerate(apps):
        y = avail_top + index * row_height
        draw.text((80, y + 13), app["application"], font=small_font, fill="#1F1F1F")
        width = int(((app["availability"] - avail_low) / (avail_high - avail_low)) * (availability_right - availability_left))
        width = max(1, min(availability_right - availability_left, width))
        draw.rounded_rectangle((availability_left, y, availability_left + width, y + 35), radius=6, fill="#70AD47")
        draw.text((availability_left + width + 12, y + 6), f"{app['availability']:.2f}%", font=small_font, fill="#1F1F1F")
    draw.line((availability_left, avail_top + len(apps) * row_height + 5, availability_right, avail_top + len(apps) * row_height + 5), fill="#555555", width=2)
    performance = FIG_DIR / "application_performance.png"
    image.save(performance)

    # Figure 4: dashboard layout
    image = Image.new("RGB", (2040, 1000), "#F9FBFE")
    draw = ImageDraw.Draw(image)
    rounded_box(draw, (45, 35, 1995, 950), "#F9FBFE", "#303030", 4, 18)
    draw.rounded_rectangle((45, 35, 1995, 155), radius=18, fill="#17365D")
    draw.rectangle((45, 120, 1995, 155), fill="#17365D")
    draw.text((110, 75), "PulseBoard APM", font=font(38, True), fill="white")
    draw.text((1530, 83), "Live monitoring overview", font=font(22), fill="#D9EAF7")
    for index, label in enumerate(["Application", "Region", "Time range"]):
        x = 120 + index * 360
        rounded_box(draw, (x, 195, x + 300, 255), "white", "#A6A6A6", 2, 10)
        center_text(draw, (x + 150, 225), label, small_font, "#555555")
    cards = [(120, "Applications", "10"), (565, "Response Time", f"{stats['response']:.0f} ms"), (1010, "CPU Usage", f"{stats['cpu']:.0f}%"), (1455, "Memory Usage", f"{stats['memory']:.0f}%")]
    for x, label, value in cards:
        rounded_box(draw, (x, 330, x + 360, 470), "white", "#D9E2F3", 3, 13)
        draw.text((x + 28, 360), label, font=small_font, fill="#666666")
        draw.text((x + 28, 398), value, font=font(32, True), fill="#17365D")
    panels = [(120, 540, 980, 860, "Performance trend chart"), (1060, 540, 1920, 860, "Application comparison chart")]
    for x1, y1, x2, y2, label in panels:
        rounded_box(draw, (x1, y1, x2, y2), "white", "#D9E2F3", 3, 13)
        draw.text((x1 + 26, y1 + 25), label, font=font(20, True), fill="#17365D")
        for n in range(8):
            x = x1 + 80 + n * 90
            h = 58 + ((n * 31) % 155)
            draw.rounded_rectangle((x, y2 - 55 - h, x + 36, y2 - 55), radius=5, fill="#5B9BD5" if n % 2 == 0 else "#70AD47")
    dashboard = FIG_DIR / "dashboard_layout.png"
    image.save(dashboard)
    return {"architecture": architecture, "composition": composition, "performance": performance, "dashboard": dashboard}


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=120) -> None:
    """Set fixed DXA geometry consistently for table, grid and every cell."""
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for index, width in enumerate(widths):
        if index < len(grid.gridCol_lst):
            grid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run(run, size=12, bold=None, italic=None, color=BLACK, font="Times New Roman", underline=None) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def set_paragraph(paragraph, before=0, after=6, line=1.15, align=None, keep_with_next=False) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if align is not None:
        paragraph.alignment = align


def add_text(doc, text: str = "", size=12, bold=False, italic=False, color=BLACK, before=0, after=6, line=1.15, align=None, keep_with_next=False, underline=None):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before, after, line, align, keep_with_next)
    if text:
        set_run(paragraph.add_run(text), size=size, bold=bold, italic=italic, color=color, underline=underline)
    return paragraph


def add_heading(doc, text: str, level: int = 1):
    sizes = {1: 16, 2: 13, 3: 12}
    paragraph = doc.add_paragraph(style=f"Report Heading {level}")
    set_paragraph(paragraph, before=14 if level == 1 else 10, after=6 if level == 1 else 4, line=1.08, keep_with_next=True)
    set_run(paragraph.add_run(text), size=sizes[level], bold=True, underline=(level == 1))
    return paragraph


def add_bullet(doc, text: str, level=0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.1
    set_run(paragraph.add_run(text), size=12)


def add_number(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.1
    set_run(paragraph.add_run(text), size=12)


def add_caption(doc, text: str) -> None:
    paragraph = add_text(doc, text, size=9, italic=True, color=MUTED, before=2, after=9, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    return paragraph


def add_page_number(footer_paragraph) -> None:
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run()
    set_run(run, size=9, color="666666")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_section_chrome(section) -> None:
    # Match the supplied academic-report format, which uses A4 portrait pages.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.70)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.30)
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph(paragraph, after=0, line=1.0)
    set_run(paragraph.add_run("Classification: Internal"), size=9, color=GREEN)
    footer = section.footer
    footer.is_linked_to_previous = False
    add_page_number(footer.paragraphs[0])


def reset_table_cell(cell) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_paragraph(paragraph, after=0, line=1.0)


def add_cover(doc, logo_path: Path | None) -> None:
    add_text(doc, "", after=34)
    add_text(doc, "PULSEBOARD: APPLICATION PERFORMANCE", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, "MONITORING (APM) DASHBOARD", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_text(doc, "[GROUP NAME]", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    add_text(doc, "BITS ZC229T: Design Project", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_text(doc, "by", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["S.No", "Student Name", "BITS ID"]
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        reset_table_cell(cell)
        shade_cell(cell, LIGHT_GREY)
        run = cell.paragraphs[0].add_run(label)
        set_run(run, size=11.5, bold=True)
    for i in range(1, 6):
        row = table.add_row().cells
        values = [str(i), f"[Student {i} - replace]", "[BITS ID]"]
        for index, value in enumerate(values):
            reset_table_cell(row[index])
            run = row[index].paragraphs[0].add_run(value)
            set_run(run, size=11)
    set_table_geometry(table, [1350, 4200, 2950])
    add_text(doc, "", after=16)
    add_text(doc, "B.Sc. Design and Computing", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_text(doc, "Design Project work carried out at", size=12.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_text(doc, "[Organisation / Location]", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    if logo_path and logo_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(logo_path), width=Inches(1.05))
        set_paragraph(paragraph, after=13)
    add_text(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_text(doc, "PILANI (RAJASTHAN)", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_text(doc, "August 2026", size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def add_submission_page(doc, logo_path: Path | None) -> None:
    add_text(doc, "", after=26)
    add_text(doc, "PULSEBOARD: APPLICATION PERFORMANCE", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, "MONITORING (APM) DASHBOARD", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=25)
    add_text(doc, "[GROUP NAME]", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_text(doc, "BITS ZC229T: Design Project", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=25)
    add_text(doc, "by", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=13)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for index, label in enumerate(["S.No", "Student Name", "BITS ID"]):
        cell = table.rows[0].cells[index]
        reset_table_cell(cell)
        shade_cell(cell, LIGHT_GREY)
        set_run(cell.paragraphs[0].add_run(label), size=11.5, bold=True)
    for i in range(1, 6):
        cells = table.add_row().cells
        for index, value in enumerate([str(i), f"[Student {i} - replace]", "[BITS ID]"]):
            reset_table_cell(cells[index])
            set_run(cells[index].paragraphs[0].add_run(value), size=11)
    set_table_geometry(table, [1350, 4200, 2950])
    add_text(doc, "", after=12)
    add_text(doc, "B.Sc. Design and Computing", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=13)
    add_text(doc, "Design Project work carried out at", size=12.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, "[Organisation / Location]", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=13)
    add_text(doc, "Submitted in partial fulfillment of the requirements for the", size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_text(doc, "B.Sc. (Design and Computing) degree programme", size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=21)
    add_text(doc, "Under the Supervision of", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_text(doc, "[Mentor Name and Designation]", size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    if logo_path and logo_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(logo_path), width=Inches(0.93))
        set_paragraph(paragraph, after=10)
    add_text(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_text(doc, "PILANI (RAJASTHAN)", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_text(doc, "August 2026", size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def add_certificate(doc) -> None:
    add_text(doc, "", after=48)
    add_text(doc, "CERTIFICATE", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    text = (
        "This is to certify that the Design Project entitled PULSEBOARD: APPLICATION PERFORMANCE "
        "MONITORING (APM) DASHBOARD and submitted by [STUDENT NAME] having BITS ID No. [BITS ID] "
        "for the partial fulfillment of the requirements of the B.Sc. (Design and Computing) degree "
        "programme of BITS embodies the bona fide work done by the student under my supervision."
    )
    add_text(doc, text, size=12.5, after=22, line=1.3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_text(doc, "", after=70)
    add_text(doc, "______________________", size=12, after=0)
    add_text(doc, "Signature of the Mentor", size=12, after=4)
    add_text(doc, "[Mentor Name]", size=12, after=2)
    add_text(doc, "[Designation / Department]", size=12, after=20)
    add_text(doc, "Place: ____________________", size=12, after=4)
    add_text(doc, "Date:  ____________________", size=12, after=0)
    add_text(doc, "Note: copy this certificate page once for each group member, if your institution requires individual certificates.", size=9.5, italic=True, color=MUTED, before=50, after=0)
    doc.add_page_break()


def add_individual_abstract(doc) -> None:
    add_text(doc, "Birla Institute of Technology & Science, Pilani", size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_text(doc, "Work-Integrated Learning Programmes Division", size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_text(doc, "BITS ZC229T: Design Project", size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, "ABSTRACT", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    details = [
        ("BITS ID No.", "[BITS ID]"),
        ("NAME OF THE STUDENT", "[Student Name]"),
        ("EMAIL ADDRESS", "[Student Email]"),
        ("STUDENT'S EMPLOYING ORGANIZATION & LOCATION", "[Organisation and Location]"),
        ("MENTOR'S NAME", "[Mentor Name]"),
        ("MENTOR'S EMAIL ADDRESS", "[Mentor Email]"),
        ("DESIGN PROJECT TITLE", "PULSEBOARD: APPLICATION PERFORMANCE MONITORING (APM) DASHBOARD"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in details:
        cells = table.add_row().cells
        for cell, content, is_label in [(cells[0], label, True), (cells[1], value, False)]:
            reset_table_cell(cell)
            if is_label:
                shade_cell(cell, LIGHT_GREY)
            set_run(cell.paragraphs[0].add_run(content), size=10.5, bold=is_label)
    set_table_geometry(table, [3350, 5150])
    add_text(doc, "", after=7)
    abstract = (
        "PulseBoard is a web-based Application Performance Monitoring dashboard developed to convert "
        "application telemetry into an accessible operational view. The system reads a structured APM "
        "dataset, derives response-time, resource, availability, error-rate and throughput metrics, and "
        "presents them through interactive filters, KPI cards, trend charts, status views and an alert table. "
        "A Node.js and Express backend exposes overview and alert endpoints; a React and Vite frontend "
        "renders the monitoring experience with reusable components and Recharts visualizations. The project "
        "demonstrates a practical path from CSV records to service-level insight. Across 14,400 observations "
        "for ten applications, the dashboard identifies a warning-level average response time (359.99 ms) and "
        "a warning-level average error rate (1.31%) while availability remains healthy at 99.74%."
    )
    p = add_text(doc, "ABSTRACT: ", size=11.5, bold=True, after=4, line=1.2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_run(p.add_run(abstract), size=11.5)
    add_text(doc, "Broad Academic Area of Work: Application Performance Monitoring, Web Development and Data Visualization", size=11, bold=True, after=7)
    add_text(doc, "Key words: APM, observability, React, Express, REST API, Recharts, CSV analytics, alerting, service health, dashboard.", size=11, bold=True, after=18)
    signatures = doc.add_table(rows=1, cols=2)
    signatures.style = "Table Grid"
    left, right = signatures.rows[0].cells
    for cell in [left, right]:
        reset_table_cell(cell)
    set_run(left.paragraphs[0].add_run("__________________\nSignature of the Student\nName: [Student Name]\nDate: [Date]\nPlace: [Place]"), size=10.5, bold=True)
    set_run(right.paragraphs[0].add_run("__________________\nSignature of the Mentor\nName: [Mentor Name]\nDate: [Date]\nPlace: [Place]"), size=10.5, bold=True)
    set_table_geometry(signatures, [4250, 4250])
    doc.add_page_break()


def add_ack_and_toc(doc) -> None:
    add_text(doc, "ACKNOWLEDGEMENTS", size=15, bold=True, underline=True, after=13)
    add_text(doc, "", after=0)
    acknowledgement = (
        "We express our sincere gratitude to our mentor, faculty members and the BITS Pilani WILP team for "
        "their guidance and encouragement during the development of PulseBoard. We also acknowledge the "
        "support of our organisation and peers, whose feedback helped refine the dashboard's usability, data "
        "presentation and operational focus."
    )
    add_text(doc, acknowledgement, size=12, after=26, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_text(doc, "TABLE OF CONTENTS", size=15, bold=True, underline=True, after=10)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, value in enumerate(["S.No", "Title", "Page No."]):
        cell = table.rows[0].cells[i]
        reset_table_cell(cell)
        shade_cell(cell, LIGHT_GREY)
        set_run(cell.paragraphs[0].add_run(value), size=10.5, bold=True)
    entries = [
        ("1", "Introduction", "7"), ("2", "Literature Review / Related Work", "8"),
        ("3", "Requirement Analysis", "9"), ("4", "System Design", "10"),
        ("5", "Implementation", "12"), ("6", "Testing and Evaluation", "15"),
        ("7", "Results and Discussion", "16"), ("8", "Inferences / Summary", "18"),
        ("9", "Conclusion and Future Work", "19"), ("10", "Reflection and Learning", "20"),
        ("11", "Bibliography", "21"), ("12", "References", "21"), ("13", "Appendices", "22"),
    ]
    for entry in entries:
        cells = table.add_row().cells
        for i, value in enumerate(entry):
            reset_table_cell(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=9.8)
            if i in (0, 2):
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [950, 6250, 1300])
    add_text(doc, "Page numbers are indicative and will update if sections are added or removed in Word.", size=9, italic=True, color=MUTED, before=5, after=0)
    doc.add_page_break()


def add_list_of_figures(doc) -> None:
    add_text(doc, "LIST OF FIGURES AND TABLES", size=15, bold=True, underline=True, after=13)
    figures = [
        "Figure 1. PulseBoard logical architecture",
        "Figure 2. Dataset composition",
        "Figure 3. Application-level performance",
        "Figure 4. PulseBoard dashboard component layout",
    ]
    tables = [
        "Table 1. Functional requirements",
        "Table 2. Non-functional requirements",
        "Table 3. Core backend modules",
        "Table 4. Threshold-based health classification",
        "Table 5. Test execution summary",
        "Table 6. Dataset-wide operational indicators",
    ]
    add_text(doc, "Figures", size=13, bold=True, after=5)
    for item in figures:
        add_number(doc, item)
    add_text(doc, "Tables", size=13, bold=True, before=12, after=5)
    for item in tables:
        add_number(doc, item)
    add_text(doc, "Source note: figures are generated from the repository's APM dataset and the implemented component structure.", size=9.5, italic=True, color=MUTED, before=18, after=0)
    doc.add_page_break()


def add_intro(doc) -> None:
    add_heading(doc, "1. INTRODUCTION", 1)
    add_heading(doc, "1.1 Background", 2)
    add_text(doc, "Modern organisations depend on many distributed applications, each of which produces operational signals such as response time, CPU utilisation, memory consumption, availability, request volume and error counts. When these signals are dispersed across files or tools, it becomes difficult for stakeholders to recognise degradation early and compare the health of services consistently. Application Performance Monitoring (APM) addresses this visibility problem by collecting, deriving and presenting health-oriented measures in a form that supports timely action.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_heading(doc, "1.2 Problem Statement", 2)
    add_text(doc, "The project addresses the challenge of turning a large set of application telemetry records into a single, understandable monitoring view. A reviewer should be able to select an application, region and reporting window, understand its current operating posture, compare services and identify potential incidents without manually calculating averages or searching raw CSV files.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_heading(doc, "1.3 Objectives", 2)
    for item in [
        "Build an interactive dashboard for summarising application performance and service health.",
        "Calculate metrics for latency, CPU, memory, availability, error rate, throughput and HTTP response counts.",
        "Classify metric values as Healthy, Warning or Critical using transparent thresholds.",
        "Provide filters and charts that support comparison across applications and regions.",
        "Expose reusable overview and alert endpoints for a React-based client.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.4 Scope", 2)
    add_text(doc, "In scope are CSV-based monitoring data, backend metric aggregation, health classification, basic alert generation, a responsive dashboard user interface and automated endpoint checks. The current implementation is a data-driven monitoring prototype: it does not ingest live telemetry streams, store historical data in a database, authenticate users or perform automated remediation.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()


def add_related_work(doc) -> None:
    add_heading(doc, "2. LITERATURE REVIEW / RELATED WORK", 1)
    add_heading(doc, "2.1 Monitoring Dashboards", 2)
    add_text(doc, "Monitoring dashboards are commonly organised around service indicators that answer practical questions: Is the system available? Are requests responding within acceptable latency? Is infrastructure under pressure? Are errors increasing? A dashboard is most useful when it combines overview metrics with drill-down information and retains the context required for a reader to decide what to investigate next. PulseBoard follows this principle by pairing KPI cards with application comparisons, time-series views, detailed summary rows and alerts.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_heading(doc, "2.2 Web Technology Context", 2)
    add_text(doc, "The frontend uses React's component model for reusable UI elements and Vite for development and production build handling. The backend uses Express to define HTTP routes and middleware around a small REST API. Recharts provides chart components for the React screen, while Papa Parse converts the source CSV into JavaScript objects. This combination is appropriate for a prototype because it separates data preparation, API logic and presentation without requiring a complex deployment platform.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_heading(doc, "2.3 Contribution of the Project", 2)
    for item in [
        "A coherent APM workflow from CSV source data to browser-based decision support.",
        "A reusable calculation layer for total metrics, per-application metrics and time-bucket summaries.",
        "Transparent health thresholds that keep alert meaning consistent across dashboard widgets.",
        "A frontend that retains a fallback sample view if the live API is temporarily unavailable.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_requirements(doc) -> None:
    add_heading(doc, "3. REQUIREMENT ANALYSIS", 1)
    add_heading(doc, "3.1 Stakeholders and User Needs", 2)
    add_text(doc, "The primary users are operations analysts, application owners and project reviewers. They need rapid visibility into overall health, the ability to compare services, evidence of performance trends and an unambiguous way to see which measures require attention. Developers additionally require a small, testable backend contract that can be consumed by the user interface.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_heading(doc, "3.2 Functional Requirements", 2)
    functional = [
        ("FR-01", "Load the provided CSV dataset and parse it into structured records."),
        ("FR-02", "Filter rows by application, region, environment, severity, event, health status and reporting window."),
        ("FR-03", "Calculate dataset-wide and per-application performance indicators."),
        ("FR-04", "Return overview data, chart data and alert data through REST endpoints."),
        ("FR-05", "Render KPI cards, filters, charts, tables and alert information in the frontend."),
        ("FR-06", "Show fallback dashboard data when the backend cannot be reached."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["ID", "Requirement"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], MID_GREY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=10.5, bold=True)
    for ident, requirement in functional:
        cells = table.add_row().cells
        for i, value in enumerate([ident, requirement]):
            reset_table_cell(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=10)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1350, 7150])
    add_caption(doc, "Table 1. Functional requirements")
    add_heading(doc, "3.3 Non-Functional Requirements", 2)
    non_functional = [
        ("Usability", "The dashboard should use clear labels, status colours and readable summaries."),
        ("Performance", "Data processing should complete quickly for the supplied 14,400-record dataset."),
        ("Reliability", "Endpoints and the build process should be verified by automated tests and a production build."),
        ("Maintainability", "Backend calculations and frontend widgets should be separated into focused modules."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["Category", "Requirement"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], MID_GREY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=10.5, bold=True)
    for category, requirement in non_functional:
        cells = table.add_row().cells
        for i, value in enumerate([category, requirement]):
            reset_table_cell(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=10)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2000, 6500])
    add_caption(doc, "Table 2. Non-functional requirements")
    doc.add_page_break()


def add_design(doc, figures: dict[str, Path]) -> None:
    add_heading(doc, "4. SYSTEM DESIGN", 1)
    add_heading(doc, "4.1 Architecture", 2)
    add_text(doc, "PulseBoard uses a lightweight three-part architecture. A CSV file is the data source; the Express backend reads and filters records, calculates aggregate values and returns JSON; the React frontend requests this data and renders it through reusable components. This design keeps metric logic on the server while allowing the client to focus on interaction and data presentation.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(figures["architecture"]), width=Inches(6.45))
    set_paragraph(para, before=4, after=2)
    add_caption(doc, "Figure 1. PulseBoard logical architecture")
    add_heading(doc, "4.2 Core Modules", 2)
    modules = [
        ("CSV Service", "Reads the local data file, parses CSV rows and applies filters."),
        ("Metrics Service", "Calculates averages, totals, p95 latency, per-application metrics and time buckets."),
        ("Health Service", "Maps calculated metric values to health labels using threshold rules."),
        ("Alert Service", "Produces concise alerts where a calculated health state needs attention."),
        ("React Components", "Render filters, KPI cards, charts, tables, status views and navigation."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["Module", "Responsibility"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], MID_GREY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=10.5, bold=True)
    for module, responsibility in modules:
        cells = table.add_row().cells
        for i, value in enumerate([module, responsibility]):
            reset_table_cell(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=10)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2450, 6050])
    add_caption(doc, "Table 3. Core backend and frontend modules")
    doc.add_page_break()

    add_heading(doc, "4.3 Data Design", 2)
    add_text(doc, "Each dataset record captures a timestamp, application, environment, region, performance values, HTTP counts, resource measures, event metadata and a health status. The current dataset covers 14,400 records from 1 May 2026 to 30 May 2026. It includes ten applications, three regions and both production and UAT environments. Numeric fields are converted to numbers before aggregation, while categorical fields form the basis of filter choices and frequency analysis.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    data_entities = [
        ("Application telemetry", "Timestamp, application, environment, region and event context."),
        ("Performance measures", "Response time, p95 latency, throughput, availability and error rate."),
        ("Resource measures", "CPU usage, memory usage and active users."),
        ("HTTP diagnostics", "4xx and 5xx response counts used to give error context."),
        ("Health output", "Healthy, Warning or Critical labels derived from defined threshold rules."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["Entity / Group", "Key Content"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], MID_GREY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=10.5, bold=True)
    for entity, content in data_entities:
        cells = table.add_row().cells
        for i, value in enumerate([entity, content]):
            reset_table_cell(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=10)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2500, 6000])
    add_caption(doc, "Data groups used by the monitoring model")
    add_heading(doc, "4.4 Dashboard Design", 2)
    add_text(doc, "The dashboard is organised as a short path from context to action: filters establish the slice of data; a hero section explains the active focus; KPI cards show primary measures; charts compare performance; and tables list summary results and alerts. This arrangement lets a reader scan the most important state before moving to detailed interpretation.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(figures["dashboard"]), width=Inches(6.35))
    set_paragraph(para, before=4, after=2)
    add_caption(doc, "Figure 4. PulseBoard dashboard component layout (derived from the implemented UI)")
    doc.add_page_break()


def add_implementation(doc, stats: dict) -> None:
    add_heading(doc, "5. IMPLEMENTATION", 1)
    add_heading(doc, "5.1 Technologies Used", 2)
    technologies = [
        ("React", "Component-based frontend rendering and state handling."),
        ("Vite", "Frontend development server and production build pipeline."),
        ("Node.js + Express", "Backend runtime and HTTP REST endpoints."),
        ("Papa Parse", "CSV parsing with header and dynamic typing options."),
        ("Recharts", "Composable charting elements for the React dashboard."),
        ("Jest + Supertest", "Automated API testing for the backend."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["Technology", "Role in PulseBoard"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], MID_GREY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=10.5, bold=True)
    for tech, role in technologies:
        cells = table.add_row().cells
        for i, value in enumerate([tech, role]):
            reset_table_cell(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=10)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2300, 6200])
    add_caption(doc, "Technology stack used in the implementation")
    add_heading(doc, "5.2 Backend API", 2)
    add_text(doc, "The Express server exposes a root health-style route, an overview route and an alerts route. The overview endpoint loads the CSV, builds the unique filter values, applies query-based selections and returns metrics, health indicators, chart data and alerts in one payload. The alerts endpoint returns the active alert list with timestamps. The API is deliberately compact so the frontend can obtain the data required for a monitoring view without joining several endpoints.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    api = [
        ("GET /", "Returns the project name and backend running status."),
        ("GET /api/overview", "Returns filters, overall metrics, health, alerts, application metrics and time-series data."),
        ("GET /api/alerts", "Returns the derived active alerts with time metadata."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["Endpoint", "Purpose"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], SOFT_GREEN)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=10.5, bold=True)
    for endpoint, purpose in api:
        cells = table.add_row().cells
        for i, value in enumerate([endpoint, purpose]):
            reset_table_cell(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=10)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2700, 5800])
    add_caption(doc, "PulseBoard REST endpoints")
    doc.add_page_break()

    add_heading(doc, "5.3 Metric and Health Logic", 2)
    add_text(doc, "Metric values are calculated from the currently filtered record set. The implementation calculates averages for response time, CPU, memory, availability, error rate and throughput; sums HTTP 4xx and 5xx counts; and derives p95 latency from the sorted latency values. The health service applies fixed thresholds. The same labels are reused by the KPI cards, overview data and alert logic, which prevents different widgets from interpreting the same value differently.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    thresholds = [
        ("CPU usage", "Healthy < 70%; Warning 70-90%; Critical > 90%"),
        ("Memory usage", "Healthy < 75%; Warning 75-90%; Critical > 90%"),
        ("Response time", "Healthy < 300 ms; Warning 300-500 ms; Critical > 500 ms"),
        ("Availability", "Healthy >= 99%; Warning 98-99%; Critical < 98%"),
        ("Error rate", "Healthy <= 1%; Warning > 1-5%; Critical > 5%"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["Measure", "Health Classification"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], SOFT_YELLOW)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=10.5, bold=True)
    for measure, classification in thresholds:
        cells = table.add_row().cells
        for i, value in enumerate([measure, classification]):
            reset_table_cell(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=10)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2300, 6200])
    add_caption(doc, "Table 4. Threshold-based health classification")
    add_heading(doc, "5.4 Frontend Behaviour", 2)
    add_text(doc, "The React application maintains selections for application, region and time range. When the selection changes, it requests a refreshed overview payload. Data is transformed into card values, summary rows, chart data and ordered alerts. The dashboard also includes a fallback data object. This makes the screen demonstrable if the deployed backend is unavailable, while an explicit status message differentiates sample data from a live API response.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_heading(doc, "5.5 Dataset Profile", 2)
    add_text(doc, f"The supplied dataset contains {stats['records']:,} records representing {stats['applications']} applications across {len(stats['regions'])} regions. The majority of records are from Production ({stats['environments'].get('Production', 0):,}), with a smaller UAT subset ({stats['environments'].get('UAT', 0):,}). The data is sufficient to show aggregate and per-application comparisons without introducing a database dependency.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()


def add_testing(doc) -> None:
    add_heading(doc, "6. TESTING AND EVALUATION", 1)
    add_heading(doc, "6.1 Test Strategy", 2)
    add_text(doc, "Testing focused on the server contract and the frontend production build. Backend unit/integration tests use Jest and Supertest to request the routes from the Express application directly. The build step runs Vite's production bundle process, confirming that the frontend components and imports resolve correctly. This provides repeatable evidence for core functionality, although it does not replace manual usability testing on a deployed system.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    tests = [
        ("T-01", "Root endpoint", "Project status JSON is returned with HTTP 200.", "Passed"),
        ("T-02", "Overview endpoint", "Metrics, health and alert fields are present with HTTP 200.", "Passed"),
        ("T-03", "Alerts endpoint", "An alerts array is returned with HTTP 200.", "Passed"),
        ("T-04", "Frontend build", "Vite production build completes without module or compilation errors.", "Passed"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, label in enumerate(["ID", "Test Area", "Expected Result", "Status"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], MID_GREY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=9.8, bold=True)
    for ident, area, expected, status in tests:
        cells = table.add_row().cells
        values = [ident, area, expected, status]
        for i, value in enumerate(values):
            reset_table_cell(cells[i])
            if i == 3:
                shade_cell(cells[i], SOFT_GREEN)
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(cells[i].paragraphs[0].add_run(value), size=9.4, bold=(i == 3))
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [900, 1700, 4650, 1250])
    add_caption(doc, "Table 5. Test execution summary")
    add_heading(doc, "6.2 Test Result", 2)
    add_text(doc, "The CI-style local verification completed successfully: three backend tests passed and the React application built successfully for production. The generated build included the HTML entry point, a CSS asset and a JavaScript bundle, demonstrating that the current codebase compiles and packages as expected.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_heading(doc, "6.3 Evaluation Limits", 2)
    add_text(doc, "The reported result verifies the application in a local development context using a static dataset. Before production use, the system should be evaluated against real-time traffic, large data volumes, network latency, authentication requirements and alert-delivery integrations. A database-backed history and observability instrumentation would also be needed for longitudinal analysis.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()


def add_results(doc, stats: dict, figures: dict[str, Path]) -> None:
    add_heading(doc, "7. RESULTS AND DISCUSSION", 1)
    add_heading(doc, "7.1 Dataset-Wide Operational Indicators", 2)
    indicators = [
        ("Average response time", f"{stats['response']:.2f} ms", "Warning (>= 300 ms)"),
        ("P95 latency", f"{stats['p95']:.0f} ms", "Tail latency needs investigation"),
        ("Average CPU usage", f"{stats['cpu']:.2f}%", "Healthy"),
        ("Average memory usage", f"{stats['memory']:.2f}%", "Healthy"),
        ("Average availability", f"{stats['availability']:.2f}%", "Healthy (>= 99%)"),
        ("Average error rate", f"{stats['error']:.2f}%", "Warning (> 1%)"),
        ("Average throughput", f"{stats['throughput']:.2f} rpm", "Contextual demand measure"),
        ("HTTP 4xx / 5xx counts", f"{stats['http4xx']:,} / {stats['http5xx']:,}", "Diagnostic volume"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, label in enumerate(["Indicator", "Observed Value", "Interpretation"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], MID_GREY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=9.8, bold=True)
    for indicator, value, interpretation in indicators:
        cells = table.add_row().cells
        for i, content in enumerate([indicator, value, interpretation]):
            reset_table_cell(cells[i])
            if "Warning" in content:
                shade_cell(cells[i], SOFT_YELLOW)
            set_run(cells[i].paragraphs[0].add_run(content), size=9.4, bold=(i == 1))
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2550, 2200, 3800])
    add_caption(doc, "Table 6. Dataset-wide operational indicators")
    add_text(doc, "The aggregate values classify CPU, memory and availability as healthy. However, the response-time average and error-rate average fall into the dashboard's warning band. Therefore, the dashboard's most useful operational function in this dataset is to direct attention toward latency and error contributors while showing that resource utilisation is not the immediate broad constraint.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(figures["composition"]), width=Inches(6.45))
    set_paragraph(para, before=4, after=2)
    add_caption(doc, "Figure 2. Dataset composition")
    doc.add_page_break()

    add_heading(doc, "7.2 Application-Level Comparison", 2)
    add_text(doc, "All ten applications report warning-level average response times under the configured threshold policy. Flipkart has the highest average response time at 412.18 ms, followed by Payment Gateway at 407.22 ms. The Payment Gateway's role in transaction completion makes it a high-priority candidate for closer review, despite the dataset-wide availability remaining healthy. Citrix Workspace has the highest p95 latency at 1,329 ms, suggesting pronounced slow-tail behaviour for at least some requests.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(figures["performance"]), width=Inches(6.45))
    set_paragraph(para, before=4, after=2)
    add_caption(doc, "Figure 3. Application-level performance")
    add_heading(doc, "7.3 Discussion", 2)
    add_text(doc, "The results illustrate why a single average should not be treated as the entire story. The average response time establishes a broad warning state, but the application comparison identifies the services that are slower relative to their peers. Availability values form a narrow healthy range, while latency and error measures carry more discriminating information. The dashboard allows these signals to be viewed together so that an analyst can form a triage hypothesis before examining logs, traces or external dependencies.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()


def add_summary_conclusion_reflection(doc) -> None:
    add_heading(doc, "8. INFERENCES / SUMMARY", 1)
    for item in [
        "The supplied data is broad enough to compare ten services while preserving region, environment and event context.",
        "Healthy average CPU, memory and availability do not eliminate a user-impacting latency or error issue.",
        "Response-time and error-rate thresholds mark the overall dataset as Warning, making these the initial investigation priorities.",
        "A component-based frontend and a modular calculation layer make the prototype easier to adapt for additional filters and live sources.",
        "The current dashboard is suitable for explanatory monitoring and data exploration; production operations would require persistent storage, access control and live ingestion.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_heading(doc, "9. CONCLUSION AND FUTURE WORK", 1)
    add_heading(doc, "9.1 Conclusion", 2)
    add_text(doc, "PulseBoard demonstrates an end-to-end APM dashboard workflow using a real project dataset and a contemporary web stack. It reads monitoring records, computes performance and health information, exposes a compact API and renders an interactive view for filtering, comparison and alert awareness. Local tests and a production frontend build both complete successfully. The project meets its aim of making telemetry-derived operational insight more accessible than raw CSV inspection.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_heading(doc, "9.2 Future Work", 2)
    for item in [
        "Ingest real-time metrics through agents, message queues or observability APIs.",
        "Store telemetry in a time-series database and apply timestamp-aware range filtering.",
        "Add authentication, role-based access and deployment-safe environment configuration.",
        "Integrate logs and distributed traces to support root-cause analysis from an alert.",
        "Add alert acknowledgement, notification channels and incident history.",
        "Extend the test suite with filter-specific, error-path and end-to-end browser tests.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_heading(doc, "10. REFLECTION AND LEARNING", 1)
    add_heading(doc, "10.1 Key Learnings", 2)
    for item in [
        "A clear metric contract makes a dashboard more reliable because calculations and health labels are shared across views.",
        "Data visualisation should support a monitoring question, not merely display available fields.",
        "Separating CSV parsing, metric aggregation, health classification and alert creation improves maintainability.",
        "Fallback UI states are valuable for demos, but must clearly disclose when data is not live.",
        "Automated endpoint testing and a production build test provide a useful baseline before deployment.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "10.2 Challenges", 2)
    add_text(doc, "The project required aligning diverse performance measures onto a simple three-level health model without masking important differences. Another practical challenge was converting a static CSV data source into an interface that behaves like a monitoring product while accurately signalling the boundary between sample/fallback information and a reachable backend API.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_heading(doc, "10.3 Improvements", 2)
    add_text(doc, "The next iteration should prioritise true timestamp parsing, a persistent data store and live telemetry ingestion. These changes would make time-range filtering semantically precise and support longitudinal health reporting. They would also enable better alert history, root-cause navigation and production-level reliability testing.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()


def add_sources_and_appendix(doc) -> None:
    add_heading(doc, "11. BIBLIOGRAPHY", 1)
    bibliography = [
        "PulseBoard Project Source Code and APM Dataset. Provided project repository, accessed August 2026.",
        "React Documentation. React Team. https://react.dev/",
        "Vite Documentation - Getting Started. VoidZero. https://vite.dev/guide/",
        "Express.js Documentation. OpenJS Foundation and Express contributors. https://expressjs.com/",
        "Papa Parse Documentation. https://www.papaparse.com/docs",
        "Recharts Documentation. https://recharts.github.io/",
        "Jest Documentation. https://jestjs.io/",
        "SuperTest Repository Documentation. https://github.com/forwardemail/supertest",
    ]
    for item in bibliography:
        add_number(doc, item)
    add_heading(doc, "12. REFERENCES", 1)
    add_text(doc, "All implementation claims in this report are based on the supplied PulseBoard repository. Technology descriptions are supported by the official documentation listed above. Dataset-wide figures and conclusions were calculated from backend/data/PulseBoard_APM_14400_Dataset_With_HTTP_Status.csv using the same aggregation approach as the project backend.", size=12, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()

    add_heading(doc, "13. APPENDICES", 1)
    add_heading(doc, "Appendix A. Project Structure", 2)
    structure = [
        ("backend/server.js", "Express application and API endpoints."),
        ("backend/services/csvService.js", "Dataset parsing and filter functions."),
        ("backend/services/metricsService.js", "Metric, p95 and application-level calculations."),
        ("backend/services/healthService.js", "Health classification wrapper."),
        ("backend/services/alertService.js", "Threshold-based alert generation."),
        ("frontend/src/App.jsx", "Application state, data fetching and dashboard assembly."),
        ("frontend/src/components", "Reusable filters, cards, charts, tables and navigation."),
        ("backend/__tests__/overview.test.js", "API checks executed through Jest and Supertest."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["Path", "Purpose"]):
        reset_table_cell(table.rows[0].cells[i])
        shade_cell(table.rows[0].cells[i], MID_GREY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(label), size=10.5, bold=True)
    for path, purpose in structure:
        cells = table.add_row().cells
        for i, value in enumerate([path, purpose]):
            reset_table_cell(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=9.6)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [3250, 5250])
    add_caption(doc, "Appendix table A1. Important source files")
    add_heading(doc, "Appendix B. Final Submission Checklist", 2)
    for item in [
        "Replace all bracketed student, ID, mentor, organisation, place and date fields.",
        "Add or duplicate individual certificate and abstract pages for every group member as required.",
        "Confirm the institution name, course code and classification statement against your official template.",
        "Update the Table of Contents page numbers after any final edits in Word.",
        "Obtain mentor and student signatures before formal submission.",
    ]:
        add_bullet(doc, item)


def crop_logo() -> Path | None:
    if not REFERENCE_PAGE.exists():
        return None
    try:
        from PIL import Image

        image = Image.open(REFERENCE_PAGE)
        # Crop the institute seal from the user-supplied sample cover page.
        crop = image.crop((520, 775, 675, 935))
        logo = FIG_DIR / "institute_logo_reference.png"
        crop.save(logo)
        return logo
    except Exception:
        return None


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for level in (1, 2, 3):
        style = doc.styles.add_style(f"Report Heading {level}", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt({1: 16, 2: 13, 3: 12}[level])
        style.font.bold = True


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = load_rows()
    stats = calculate_stats(rows)
    figures = save_figures(stats)
    logo_path = crop_logo()
    doc = Document()
    configure_styles(doc)
    set_section_chrome(doc.sections[0])
    add_cover(doc, logo_path)
    add_submission_page(doc, logo_path)
    add_certificate(doc)
    add_individual_abstract(doc)
    add_ack_and_toc(doc)
    add_list_of_figures(doc)
    add_intro(doc)
    add_related_work(doc)
    add_requirements(doc)
    add_design(doc, figures)
    add_implementation(doc, stats)
    add_testing(doc)
    add_results(doc, stats, figures)
    add_summary_conclusion_reflection(doc)
    add_sources_and_appendix(doc)
    doc.core_properties.title = "PulseBoard APM Design Project Report"
    doc.core_properties.subject = "BITS ZC229T Design Project"
    doc.core_properties.author = "[Student / Group Name]"
    doc.core_properties.comments = "Generated from the PulseBoard project repository."
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
